from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE = next(ROOT.glob("1 *.docx"))
OUT = ROOT / "Chuyen de - Ung dung Deep Learning cho anh xa dieu che trong he thong BIBCM-ID.docx"
IMG = ROOT / "images"


TITLE = "ỨNG DỤNG DEEP LEARNING CHO ÁNH XẠ ĐIỀU CHẾ TRONG HỆ THỐNG BIBCM-ID"
AUTHOR = "Lại Tiến Đệ"
MAJOR = "Kỹ thuật viễn thông"
SUPERVISOR = "TS. Phạm Xuân Nghĩa"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_base_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 13, "2F5597"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def clear_header_footer_parts(doc: Document) -> None:
    for part in doc.part.package.parts:
        name = str(part.partname)
        if not (name.startswith("/word/header") or name.startswith("/word/footer")):
            continue
        root = part.element
        for child in list(root):
            root.remove(child)
        root.append(OxmlElement("w:p"))


def add_centered(doc, text, size=13, bold=False, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(spacing_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    return p


def add_para(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(13)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Times New Roman"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r2.font.size = Pt(13)
        r2.italic = italic
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(13)
        r.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    p.add_run("• ").bold = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(13)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(12)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)


def add_image(doc, rel_path, caption, width_cm=13.5):
    path = IMG / rel_path
    if not path.exists():
        add_para(doc, f"[Không tìm thấy hình: {rel_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_small_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for style_name in ("Table Grid", "TableGrid", "Light Grid"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        set_cell_shading(c, "D9EAF7")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(c, widths[i])
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                r.font.size = Pt(12)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(12)
    doc.add_paragraph()


def build_doc():
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    apply_base_styles(doc)

    add_centered(doc, "HỌC VIỆN KỸ THUẬT QUÂN SỰ", 13, True, 0)
    add_centered(doc, "KHOA VÔ TUYẾN ĐIỆN TỬ", 13, True, 36)
    add_centered(doc, "BÁO CÁO CHUYÊN ĐỀ NGHIÊN CỨU", 18, True, 36)
    add_centered(doc, "TÊN CHUYÊN ĐỀ:", 14, True, 8)
    add_centered(doc, TITLE, 17, True, 48)
    add_centered(doc, f"HỌC VIÊN THỰC HIỆN: {AUTHOR}", 13, True, 6)
    add_centered(doc, f"CHUYÊN NGÀNH: {MAJOR}", 13, True, 6)
    add_centered(doc, f"CÁN BỘ HƯỚNG DẪN: {SUPERVISOR}", 13, True, 96)
    add_centered(doc, "HÀ NỘI - 2026", 13, True, 0)
    doc.add_page_break()

    doc.add_heading("MỤC LỤC", level=1)
    toc = [
        "I. MỞ ĐẦU",
        "II. CƠ SỞ ÁNH XẠ ĐIỀU CHẾ VÀ HỆ THỐNG BIBCM-ID",
        "III. MÔ HÌNH AUTOENCODER CHO ÁNH XẠ ĐIỀU CHẾ",
        "IV. HUẤN LUYỆN, MÔ PHỎNG VÀ ĐÁNH GIÁ HIỆU NĂNG",
        "V. KẾT LUẬN",
        "VI. TÀI LIỆU THAM KHẢO",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    doc.add_heading("I. MỞ ĐẦU", level=1)
    add_para(doc, "Trong hệ thống truyền thông số hiện đại, ánh xạ điều chế là khâu chuyển đổi chuỗi bit rời rạc thành điểm tín hiệu trên mặt phẳng IQ để truyền qua kênh vật lý. Với các hệ thống dùng điều chế bậc cao như M-PSK hoặc M-QAM, ánh xạ này quyết định trực tiếp khoảng cách Euclid tối thiểu giữa các ký hiệu, độ nhạy với nhiễu pha, méo biên độ và khả năng khai thác thông tin mềm tại bộ thu.")
    add_para(doc, "BIBCM-ID (Bit-Interleaved Block Coded Modulation with Iterative Decoding) kết hợp mã kênh, hoán vị bit, điều chế bậc cao và giải mã lặp. Hiệu năng của hệ thống phụ thuộc không chỉ vào bộ giải mã LDPC/SISO mà còn phụ thuộc mạnh vào cấu trúc ánh xạ bit-ký hiệu. Khi kênh có fading Rayleigh, sai lệch tần số sóng mang (CFO), méo khuếch đại công suất (PA) và nhiễu liên ký hiệu (ISI), ánh xạ điều chế cố định kiểu 16-QAM thường không còn tối ưu.")
    add_para(doc, "Deep Learning mở ra hướng tiếp cận mới: xem bộ phát, kênh truyền và bộ thu như một AutoEncoder khả vi đầu-cuối. Trong mô hình này, mạng Encoder học ánh xạ điều chế thích nghi, kênh truyền đóng vai trò lớp ngẫu nhiên hoặc lớp mô phỏng vật lý, còn Decoder học quá trình giải điều chế và cân bằng kênh. Chuyên đề này tập trung trình bày việc ứng dụng Deep Learning cho ánh xạ điều chế trong hệ thống BIBCM-ID, dựa trên nội dung chương 3 của luận văn và diễn giải theo hướng chuyên đề nghiên cứu.")
    add_para(doc, "Ngoài phần mở đầu và kết luận, chuyên đề gồm ba nội dung chính:")
    add_bullet(doc, "Cơ sở ánh xạ điều chế, các suy hao kênh không lý tưởng và mối liên hệ với BIBCM-ID.")
    add_bullet(doc, "Kiến trúc AutoEncoder thay thế khối ánh xạ/giải ánh xạ điều chế trong chuỗi truyền thông.")
    add_bullet(doc, "Phương pháp huấn luyện, mô phỏng và đánh giá hiệu năng so với 16-QAM truyền thống.")

    doc.add_heading("II. CƠ SỞ ÁNH XẠ ĐIỀU CHẾ VÀ HỆ THỐNG BIBCM-ID", level=1)
    doc.add_heading("2.1. Nguyên lý ánh xạ điều chế số", level=2)
    add_para(doc, "Điều chế số biến chuỗi bit nhị phân thành các ký hiệu phức thích hợp với kênh truyền. Tại bộ điều chế, mỗi nhóm k = log2(M) bit được ánh xạ tới một điểm trong chòm sao M mức trên mặt phẳng đồng pha và vuông pha. Tín hiệu băng gốc của ký hiệu thứ n có thể viết:")
    add_equation(doc, "x[n] = aI[n] + j · aQ[n]")
    add_para(doc, "Trong đó aI[n] và aQ[n] là thành phần đồng pha và vuông pha. Với M-PSK, thông tin chủ yếu nằm ở pha tín hiệu, còn biên độ gần như không đổi. Với M-QAM, thông tin được mang đồng thời bởi biên độ và pha. Khi bậc điều chế M tăng, hiệu suất phổ được cải thiện nhưng khoảng cách Euclid tối thiểu dmin giữa các điểm giảm, làm xác suất lỗi ký hiệu tăng trong điều kiện nhiễu.")
    add_para(doc, "Ở bộ thu, quyết định cứng thông thường chọn điểm chòm sao gần nhất với tín hiệu nhận được:")
    add_equation(doc, "m̂ = arg min || y − xm ||²,  m ∈ {0, ..., M − 1}")
    add_para(doc, "Trong BIBCM-ID, bộ giải điều chế không chỉ tạo quyết định cứng mà còn sinh thông tin mềm dạng LLR để trao đổi lặp với bộ giải mã. Vì vậy, một ánh xạ tốt cần đồng thời tạo khoảng cách hình học thuận lợi và hỗ trợ lan truyền thông tin ngoại lai ổn định qua các vòng lặp giải mã.")

    doc.add_heading("2.2. Các suy hao kênh không lý tưởng", level=2)
    add_para(doc, "Trong triển khai thực tế, tín hiệu vô tuyến không chỉ chịu nhiễu AWGN mà còn chịu nhiều dạng suy hao đồng thời. Ba nhóm suy hao quan trọng trong chuyên đề gồm méo phi tuyến của bộ khuếch đại công suất, sai lệch tần số sóng mang và fading đa đường.")
    add_para(doc, "Méo PA xuất hiện khi bộ khuếch đại làm việc gần vùng bão hòa để tăng hiệu suất năng lượng. Mô hình Rapp mô tả đáp ứng AM/AM và AM/PM bằng biểu thức:")
    add_equation(doc, "g(r) = r / (1 + (r / Asat)^(2p))^(1/(2p)),   Φ(r) = αr² / (1 + βr²)")
    add_para(doc, "Với 16-QAM, các điểm ở biên chòm sao dễ bị nén biên độ hơn các điểm gần gốc tọa độ, làm biến dạng lưới chòm sao và giảm dmin. Điều này gây tăng BER, đặc biệt ở vùng SNR cao nơi lỗi do nhiễu giảm nhưng lỗi do méo phần cứng vẫn còn.")
    add_image(doc, "PA characteristic curve 1.png", "Hình 1. Đặc tuyến biên độ của bộ khuếch đại công suất theo mô hình Rapp.", 13.2)
    add_image(doc, "16QAM-constellation-afterPA.png", "Hình 2. Chòm sao 16-QAM sau khi chịu méo phi tuyến PA.", 10.5)
    add_para(doc, "Sai lệch tần số sóng mang (CFO) do không đồng bộ dao động giữa máy phát và máy thu gây quay pha tích lũy theo thời gian:")
    add_equation(doc, "ycfo[n] = x[n] · exp(j2πΔfTs n)")
    add_para(doc, "Góc quay tăng tuyến tính theo chỉ số ký hiệu n. Với M-PSK và M-QAM, hiện tượng này làm điểm nhận được trượt khỏi vùng quyết định, khiến các bộ giải điều chế dựa trên chòm sao cố định phải dùng thêm PLL hoặc khối bù CFO. Cách xử lý tách rời này thường không tối ưu khi đồng thời tồn tại PA, fading và nhiễu.")
    add_image(doc, "CFO-effect.png", "Hình 3. CFO gây quay pha tiến triển trên chuỗi điểm chòm sao.", 13.2)
    add_para(doc, "Fading Rayleigh đa đường làm tín hiệu đến bộ thu qua nhiều đường phản xạ có độ trễ và biên độ khác nhau. Mô hình kênh có thể viết:")
    add_equation(doc, "y[n] = Σ h[l] · x[n − l] + w[n]")
    add_para(doc, "Trong đó h[l] là hệ số kênh phức phân bố Rayleigh, còn w[n] là nhiễu AWGN. Khi kết hợp với BIBCM-ID, sai lệch ở khối giải điều chế sẽ ảnh hưởng trực tiếp đến LLR đầu vào của bộ giải mã, làm giảm tốc độ hội tụ của quá trình giải mã lặp.")

    doc.add_heading("2.3. Vai trò của ánh xạ điều chế trong BIBCM-ID", level=2)
    add_para(doc, "BIBCM-ID sử dụng hoán vị bit trước khi ánh xạ lên chòm sao, sau đó bộ thu trao đổi thông tin ngoại lai giữa bộ giải điều chế và bộ giải mã. Trong mỗi vòng lặp, thông tin a priori từ bộ giải mã giúp bộ giải điều chế tính LLR chính xác hơn cho từng bit trong ký hiệu, còn thông tin ngoại lai từ bộ giải điều chế lại cải thiện quyết định của bộ giải mã.")
    add_para(doc, "Do đó, ánh xạ điều chế không chỉ là bảng tra bit-ký hiệu. Nó xác định cách các bit trong cùng một ký hiệu tương tác, ảnh hưởng tới độ độc lập của các kênh bit tương đương và khả năng giảm lan truyền lỗi trong quá trình lặp. Khi ánh xạ được học bằng Deep Learning, hệ thống có thể tự điều chỉnh hình học chòm sao để phù hợp với đặc tính kênh và mục tiêu tối ưu BER.")

    doc.add_heading("III. MÔ HÌNH AUTOENCODER CHO ÁNH XẠ ĐIỀU CHẾ", level=1)
    doc.add_heading("3.1. Ý tưởng mô hình hóa đầu-cuối", level=2)
    add_para(doc, "Một liên kết truyền thông có thể được biểu diễn như một AutoEncoder: Encoder đóng vai trò bộ phát, kênh truyền là lớp biến đổi ngẫu nhiên hoặc lớp mô phỏng suy hao vật lý, còn Decoder là bộ thu. Toàn bộ tham số được tối ưu chung bằng thuật toán lan truyền ngược theo một hàm mất mát duy nhất.")
    add_image(doc, "autoencoder.png", "Hình 4. Kiến trúc AutoEncoder cơ bản cho hệ thống truyền thông.", 12.0)
    add_para(doc, "Khác với thiết kế truyền thống vốn tối ưu riêng từng khối như điều chế, cân bằng, bù CFO và giải điều chế, AutoEncoder học trực tiếp ánh xạ từ thông điệp đầu vào tới xác suất khôi phục thông điệp ở đầu ra. Cách tiếp cận này đặc biệt phù hợp với BIBCM-ID vì nó có thể tối ưu cả hình học chòm sao và độ tin cậy thông tin mềm.")
    add_image(doc, "sys_diagram.png", "Hình 5. Sơ đồ hệ thống đầu-cuối khả vi dùng trong mô phỏng.", 14.2)

    doc.add_heading("3.2. Kiến trúc một ký hiệu", level=2)
    add_para(doc, "Kiến trúc thứ nhất thay thế trực tiếp bộ điều chế và giải điều chế 16-QAM bằng mạng nơ-ron. Với M = 16, mỗi nhóm k = 4 bit được biểu diễn bởi vector one-hot s có 16 chiều. Encoder gồm các lớp fully-connected tạo đầu ra hai chiều [xI, xQ], sau đó qua lớp chuẩn hóa công suất để bảo đảm năng lượng phát trung bình không vượt chuẩn.")
    add_para(doc, "Kênh truyền trong mô phỏng bao gồm Rayleigh fading, nhiễu AWGN và méo PA. Decoder nhận vector [yI, yQ, hI, hQ, SNR], qua các lớp fully-connected kết hợp Batch Normalization và ReLU, sau cùng là lớp Softmax 16 lớp. Nhờ biết CSI và mức SNR, Decoder có thể học đồng thời giải điều chế và cân bằng mà không cần bộ ZF/MMSE tách rời.")
    add_para(doc, "Kết quả học được không nhất thiết là lưới vuông như 16-QAM. Các điểm chòm sao có thể dịch chuyển, co giãn hoặc phân bố không đều để giảm tác động của PA và fading, miễn là xác suất giải mã đúng sau kênh được tối đa hóa.")

    doc.add_heading("3.3. Kiến trúc nhiều ký hiệu và bù CFO", level=2)
    add_para(doc, "Trong trường hợp có CFO, suy hao phụ thuộc theo thời gian nên một ký hiệu đơn lẻ không đủ mô tả quỹ đạo quay pha. Kiến trúc nhiều ký hiệu dùng Encoder chia sẻ theo chuỗi thời gian để ánh xạ một block thông tin thành chuỗi điểm IQ. Decoder xử lý toàn bộ chuỗi, từ đó học được quy luật quay pha tiến triển mà không cần pilot hoặc PLL rõ ràng.")
    add_image(doc, "multisymbol Auto Encoder.png", "Hình 6. Kiến trúc TimeDistributed Encoder cho mô hình nhiều ký hiệu.", 12.5)
    add_para(doc, "Điểm quan trọng của mô hình là khả năng học bù CFO gián tiếp. Thay vì ước lượng Δf rồi bù pha theo công thức cố định, mạng học biểu diễn thống kê của chuỗi nhận được và đưa ra quyết định phân loại cuối cùng. Trong BIBCM-ID, hướng này có thể giảm áp lực cho vòng lặp giải điều chế/giải mã khi kênh có sai lệch pha kéo dài.")

    doc.add_heading("3.4. Neural FEC trong ánh xạ điều chế", level=2)
    add_para(doc, "Kiến trúc thứ ba mở rộng Encoder để ánh xạ k bit đầu vào thành N ký hiệu IQ, với N lớn hơn số ký hiệu tối thiểu cần thiết. Khi đó, mạng tự tạo dư thừa giống như mã sửa sai truyền thống. Tốc độ mã hiệu dụng được xác định:")
    add_equation(doc, "R = k / (N · log2(M))")
    add_para(doc, "Cách tiếp cận Neural FEC làm mờ ranh giới giữa điều chế và mã hóa kênh. Trong BIBCM-ID, nó có thể đóng vai trò lớp ánh xạ điều chế có khả năng tự bảo vệ trước suy hao, đồng thời vẫn kết hợp với mã LDPC hoặc bộ giải mã lặp ở tầng sau. Tuy nhiên, để triển khai thực tế cần kiểm soát độ trễ, kích thước block và khả năng tổng quát hóa khi kênh thay đổi.")

    doc.add_heading("IV. HUẤN LUYỆN, MÔ PHỎNG VÀ ĐÁNH GIÁ HIỆU NĂNG", level=1)
    doc.add_heading("4.1. Phương pháp huấn luyện", level=2)
    add_para(doc, "Quá trình huấn luyện được thực hiện offline trên các mẫu sinh ngẫu nhiên. Đầu vào là chỉ số thông điệp hoặc vector one-hot; đầu ra mong muốn là chính nhãn thông điệp ban đầu. Hàm mất mát sử dụng cross-entropy giữa phân bố Softmax của Decoder và nhãn đúng.")
    add_para(doc, "Để tránh mô hình chỉ tốt trong một điều kiện hẹp, dữ liệu huấn luyện được sinh trên nhiều mức SNR và nhiều cấu hình suy hao. Curriculum learning có thể được sử dụng bằng cách bắt đầu từ kênh nhẹ, sau đó tăng dần mức CFO, PA và fading. Cách này giúp mạng ổn định trước khi học các biến dạng phức tạp.")
    add_small_table(
        doc,
        ["Kịch bản", "Mục tiêu học", "Suy hao chính", "Ý nghĩa trong BIBCM-ID"],
        [
            ["AE một ký hiệu", "Học chòm sao chống méo PA", "Rayleigh + PA + AWGN", "Cải thiện chất lượng LLR đầu vào bộ giải mã"],
            ["AE nhiều ký hiệu", "Học bù CFO theo chuỗi", "Rayleigh + CFO + AWGN", "Giảm lỗi pha tích lũy qua các vòng lặp"],
            ["Neural FEC", "Tạo dư thừa ngay trong ánh xạ", "PA + CFO + AWGN", "Kết hợp mã hóa và điều chế ở mức học sâu"],
        ],
        [2.7, 4.3, 3.2, 5.0],
    )

    doc.add_heading("4.2. Kết quả chòm sao học được", level=2)
    add_para(doc, "Kết quả mô phỏng cho thấy AutoEncoder không sao chép nguyên dạng 16-QAM mà tự học một hình học chòm sao mới. Các điểm có xu hướng phân bố để giảm tác động của nén biên độ, tránh vùng dễ bão hòa của PA và tăng khả năng phân tách sau fading. Đây là minh chứng rằng Deep Learning có thể học ánh xạ điều chế theo kênh, thay vì áp dụng bảng ánh xạ cố định.")
    add_image(doc, "AE-learned-consteallation.png", "Hình 7. Chòm sao tín hiệu do Encoder tự học.", 11.5)

    doc.add_heading("4.3. So sánh BER với 16-QAM", level=2)
    add_para(doc, "Ở kịch bản một ký hiệu, AE đạt BER thấp hơn 16-QAM khi tồn tại đồng thời Rayleigh fading, PA và AWGN. Lợi ích lớn nhất xuất hiện ở vùng SNR trung bình-cao, nơi méo phi tuyến trở thành nguồn lỗi đáng kể. So với 16-QAM truyền thống, chòm sao học được có khả năng chống nén biên độ tốt hơn và duy trì khoảng cách quyết định hiệu dụng lớn hơn.")
    add_image(doc, "Scenario 1 result-BER AE vs 16QAM .png", "Hình 8. So sánh BER giữa AE một ký hiệu và 16-QAM.", 13.2)
    add_para(doc, "Với CFO thấp và trung bình, AE nhiều ký hiệu duy trì ưu thế so với 16-QAM kể cả khi baseline có bù CFO lý tưởng. Ở mục tiêu BER khoảng 10^-3, mô hình đạt lợi ích xấp xỉ 2-3 dB trong các điều kiện mô phỏng được khảo sát. Điều này cho thấy mạng không chỉ học hình học chòm sao mà còn học quan hệ động theo thời gian của chuỗi ký hiệu.")
    add_image(doc, "Scenario 2 result-BER AE vs 16QAM CFO 0.005 .png", "Hình 9. BER tại mức CFO thấp ΔfTs = 0.005.", 13.2)
    add_image(doc, "Scenario 2 result-BER AE vs 16QAM CFO 0.01 .png", "Hình 10. BER tại mức CFO trung bình ΔfTs = 0.01.", 13.2)
    add_para(doc, "Trong kịch bản Neural FEC, khi Encoder phát nhiều ký hiệu IQ cho một nhóm bit đầu vào, hệ thống tự tạo dư thừa và cải thiện khả năng phục hồi lỗi. Đây là hướng có tiềm năng cho các hệ thống BIBCM-ID thế hệ mới, nơi mã kênh và ánh xạ điều chế có thể được đồng tối ưu thay vì thiết kế tuần tự.")
    add_image(doc, "Scenario 3 AE with channel coding feature .png", "Hình 11. Hiệu năng AE khi tích hợp chức năng mã hóa kênh.", 13.2)

    doc.add_heading("4.4. Nhận xét về khả năng ứng dụng", level=2)
    add_small_table(
        doc,
        ["Tiêu chí", "Ưu điểm của ánh xạ học sâu", "Thách thức cần xử lý"],
        [
            ["Hiệu năng BER", "Tự thích nghi với PA, CFO và fading; cải thiện vùng SNR trung bình-cao.", "Cần kiểm chứng trên kênh đo thực tế và nhiều cấu hình phần cứng."],
            ["Tích hợp BIBCM-ID", "Có thể tạo LLR chất lượng hơn cho giải mã lặp, giảm lan truyền lỗi.", "Cần thiết kế giao diện thông tin mềm ổn định với bộ giải mã LDPC/SISO."],
            ["Độ phức tạp", "Inference có thể song song hóa trên GPU/FPGA/ASIC.", "Chi phí bộ nhớ và độ trễ cần được ràng buộc khi triển khai thời gian thực."],
            ["Khả năng tổng quát", "Huấn luyện đa SNR và curriculum learning giúp tăng độ bền.", "Nguy cơ overfitting khi kênh vận hành khác miền huấn luyện."],
        ],
        [3.0, 6.0, 5.6],
    )
    add_para(doc, "Từ góc nhìn BIBCM-ID, giá trị chính của Deep Learning không nằm ở việc thay thế mọi khối xử lý truyền thống bằng một mô hình hộp đen, mà nằm ở khả năng đồng tối ưu các khối vốn phụ thuộc mạnh lẫn nhau. Ánh xạ điều chế học được có thể được thiết kế để tạo thông tin mềm tốt hơn, chống suy hao phần cứng tốt hơn và phối hợp hiệu quả hơn với giải mã lặp.")

    doc.add_heading("V. KẾT LUẬN", level=1)
    add_para(doc, "Chuyên đề đã trình bày việc ứng dụng Deep Learning cho ánh xạ điều chế trong hệ thống BIBCM-ID thông qua mô hình AutoEncoder đầu-cuối. Nội dung cho thấy ánh xạ điều chế không nên chỉ được xem là bảng bit-ký hiệu cố định, mà là một thành phần có thể học và tối ưu theo đặc tính kênh, đặc biệt khi tồn tại PA phi tuyến, CFO và fading Rayleigh.")
    add_para(doc, "Các kết quả mô phỏng trong chương 3 của luận văn cho thấy AutoEncoder có thể học chòm sao chống méo PA, bù CFO gián tiếp không cần pilot trong mô hình nhiều ký hiệu, và tích hợp chức năng Neural FEC bằng cách tạo dư thừa trong ánh xạ IQ. So với 16-QAM truyền thống, các mô hình AE đạt lợi ích rõ rệt về BER trong nhiều điều kiện kênh không lý tưởng.")
    add_para(doc, "Đối với hệ thống BIBCM-ID, hướng nghiên cứu này có ý nghĩa thực tiễn vì chất lượng ánh xạ và giải ánh xạ ảnh hưởng trực tiếp đến LLR trao đổi với bộ giải mã lặp. Trong các nghiên cứu tiếp theo, cần tập trung vào ba vấn đề: sinh LLR mềm từ Decoder một cách ổn định, đồng huấn luyện với bộ giải mã LDPC/ANN, và kiểm chứng trên kênh thực hoặc nền tảng SDR để đánh giá khả năng triển khai.")

    doc.add_heading("VI. TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "T. J. O'Shea and J. Hoydis, “An Introduction to Deep Learning for the Physical Layer,” IEEE Transactions on Cognitive Communications and Networking, 2017.",
        "S. Dörner, S. Cammerer, J. Hoydis and S. ten Brink, “Deep Learning Based Communication Over the Air,” IEEE Journal of Selected Topics in Signal Processing, 2018.",
        "F. A. Aoudia and J. Hoydis, “Model-Free Training of End-to-End Communication Systems,” IEEE Journal on Selected Areas in Communications, 2019.",
        "H. Ye, G. Y. Li and B. H. Juang, “Power of Deep Learning for Channel Estimation and Signal Detection in OFDM Systems,” IEEE Wireless Communications Letters, 2018.",
        "T. Gruber, S. Cammerer, J. Hoydis and S. ten Brink, “On Deep Learning-Based Channel Decoding,” CISS, 2017.",
        "M. Stark, F. A. Aoudia and J. Hoydis, “Joint Learning of Geometric and Probabilistic Constellation Shaping,” IEEE GLOBECOM Workshops, 2019.",
        "J. G. Proakis and M. Salehi, Digital Communications, McGraw-Hill, 2008.",
        "Nội dung chương 3, file LuanVan_De.tex: “Intelligent AutoEncoder-Based Modulation”.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(3)

    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    clear_header_footer_parts(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
